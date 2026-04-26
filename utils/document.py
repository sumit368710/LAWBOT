# utils/document_loader.py
# DEBUG VERSION - shows exactly what file is loaded and what embeddings are used

import os
import glob
import json
import re
import base64
import hashlib

print("🔥 DEBUG: document_loader.py imported")
print("🔥 DEBUG FILE PATH:", __file__)
print("🔥 DEBUG CURRENT WORKING DIR:", os.getcwd())

from groq import Groq

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS


# =====================================================
# PATHS
# =====================================================
BASE_DIR = os.path.dirname(os.path.dirname(__file__))

DOCUMENTS_FOLDER = os.path.join(BASE_DIR, "files")
FAISS_INDEX_PATH = os.path.join(BASE_DIR, "faiss_index")
METADATA_PATH = os.path.join(BASE_DIR, "file_metadata.json")

print("🔥 DEBUG BASE_DIR:", BASE_DIR)
print("🔥 DEBUG DOCUMENTS_FOLDER:", DOCUMENTS_FOLDER)
print("🔥 DEBUG FAISS_INDEX_PATH:", FAISS_INDEX_PATH)


# =====================================================
# GROQ CLIENT
# =====================================================
groq_client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)

print("🔥 DEBUG GROQ KEY FOUND:", bool(os.getenv("GROQ_API_KEY")))


# =====================================================
# FILE HASH
# =====================================================
def get_file_hash(path):
    hasher = hashlib.md5()

    with open(path, "rb") as f:
        while chunk := f.read(8192):
            hasher.update(chunk)

    return hasher.hexdigest()


# =====================================================
# CLEAN TEXT
# =====================================================
def clean_text(text):

    text = text.replace("\n", " ")
    text = text.replace("\t", " ")

    text = re.sub(r'@\w+', '', text)
    text = re.sub(r'Answer:\s*Answer:', 'Answer:', text)
    text = re.sub(r'Question:\s*Question:', 'Question:', text)
    text = re.sub(r'\s+', ' ', text)

    return text.strip()


# =====================================================
# EMBEDDINGS
# =====================================================
def get_embeddings():

    print("🔥 DEBUG: Loading HuggingFace embeddings")
    print("🔥 DEBUG MODEL: sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True}
    )

    print("✅ DEBUG: HuggingFace embeddings loaded successfully")

    return embeddings


# =====================================================
# OCR
# =====================================================
def groq_image_ocr(image_bytes):

    print("🔥 DEBUG: OCR started")

    b64 = base64.b64encode(image_bytes).decode()

    response = groq_client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Extract all text exactly from this page."
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{b64}"
                        }
                    }
                ]
            }
        ]
    )

    print("✅ DEBUG: OCR success")

    return response.choices[0].message.content


# =====================================================
# DOCUMENT LOADER
# =====================================================
class DocumentLoader:

    def __init__(self):

        print("🔥 DEBUG: DocumentLoader initialized")

        self.embeddings = get_embeddings()

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=700,
            chunk_overlap=120,
            separators=[
                "\nQ.",
                "\nQuestion:",
                "\n\n",
                "\n",
                ". ",
                " "
            ]
        )

        print("✅ DEBUG: Splitter ready")

    # =================================================
    # MAIN FUNCTION
    # =================================================
    def load_from_folder(self, folder_path=DOCUMENTS_FOLDER):

        print("🔥 DEBUG: load_from_folder started")
        print("🔥 DEBUG folder_path:", folder_path)

        os.makedirs(folder_path, exist_ok=True)

        files = list(set(
            glob.glob(os.path.join(folder_path, "**", "*.pdf"), recursive=True) +
            glob.glob(os.path.join(folder_path, "**", "*.docx"), recursive=True)
        ))

        print("🔥 DEBUG total files found:", len(files))

        if not files:
            raise FileNotFoundError("❌ No files found in files folder")

        old_metadata = {}

        if os.path.exists(METADATA_PATH):
            with open(METADATA_PATH, "r", encoding="utf-8") as f:
                old_metadata = json.load(f)

        print("🔥 DEBUG old metadata count:", len(old_metadata))

        vectorstore = None

        if os.path.exists(FAISS_INDEX_PATH):
            try:
                print("🔥 DEBUG loading old FAISS index")

                vectorstore = FAISS.load_local(
                    FAISS_INDEX_PATH,
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )

                print("✅ DEBUG FAISS loaded")

            except Exception as e:
                print("❌ DEBUG FAISS load failed:", e)
                vectorstore = None

        new_docs = []
        updated_metadata = old_metadata.copy()

        for fp in files:

            fname = os.path.basename(fp)
            file_hash = get_file_hash(fp)

            print("\n🔍 DEBUG Processing:", fname)

            if fname in old_metadata and old_metadata[fname] == file_hash:
                print("⏩ DEBUG Skipped unchanged file")
                continue

            text = self.extract_text(fp)

            if not text.strip():
                print("❌ DEBUG Empty file")
                continue

            text = clean_text(text)

            chunks = self.splitter.split_text(text)

            print("🔥 DEBUG chunks created:", len(chunks))

            docs = [
                Document(
                    page_content=chunk,
                    metadata={"source": fname}
                )
                for chunk in chunks
                if len(chunk.strip()) > 50
            ]

            new_docs.extend(docs)
            updated_metadata[fname] = file_hash

            print("✅ DEBUG docs added:", len(docs))

        # create/update index
        if vectorstore is None:

            print("🔥 DEBUG creating new FAISS")

            vectorstore = FAISS.from_documents(
                new_docs,
                self.embeddings
            )

            print("✅ DEBUG new FAISS created")

        elif new_docs:

            print("🔥 DEBUG adding new docs to FAISS")

            vectorstore.add_documents(new_docs)

            print("✅ DEBUG docs added to FAISS")

        else:
            print("⚡ DEBUG no new docs")

        vectorstore.save_local(FAISS_INDEX_PATH)

        print("✅ DEBUG FAISS saved")

        with open(METADATA_PATH, "w", encoding="utf-8") as f:
            json.dump(updated_metadata, f, indent=2)

        print("✅ DEBUG metadata saved")

        return vectorstore

    # =================================================
    # EXTRACT TEXT
    # =================================================
    def extract_text(self, path):

        if path.endswith(".pdf"):
            return self.extract_pdf(path)

        return self.extract_docx(path)

    # =================================================
    # PDF
    # =================================================
    def extract_pdf(self, path):

        try:
            import fitz

            print("🔥 DEBUG PDF extraction:", path)

            text = ""

            with fitz.open(path) as pdf:
                for page in pdf:
                    text += page.get_text()

            if text.strip():
                print("✅ DEBUG PyMuPDF success")
                return text

        except Exception as e:
            print("❌ DEBUG PyMuPDF failed:", e)

        print("🚀 DEBUG OCR fallback")

        try:
            import fitz

            pdf = fitz.open(path)
            final_text = ""

            for page in pdf:
                pix = page.get_pixmap()
                img_bytes = pix.tobytes("png")

                page_text = groq_image_ocr(img_bytes)
                final_text += page_text + "\n\n"

            return final_text

        except Exception as e:
            print("❌ DEBUG OCR failed:", e)

        return ""

    # =================================================
    # DOCX
    # =================================================
    def extract_docx(self, path):

        try:
            from docx import Document as DocxDoc

            print("🔥 DEBUG DOCX extraction:", path)

            doc = DocxDoc(path)
            texts = []

            for p in doc.paragraphs:
                if p.text.strip():
                    texts.append(p.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)

            print("✅ DEBUG DOCX extracted")

            return "\n\n".join(texts)

        except Exception as e:
            print("❌ DEBUG DOCX failed:", e)
            return ""













# ## from groq api embedding

# import os
# import glob
# import json
# import time
# import base64
# import hashlib
# from typing import List

# import cohere
# from groq import Groq

# from langchain_core.embeddings import Embeddings
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_core.documents import Document
# from langchain_community.vectorstores import FAISS

# # =========================
# # PATHS
# # =========================
# BASE_DIR = os.path.dirname(os.path.dirname(__file__))
# DOCUMENTS_FOLDER = os.path.join(BASE_DIR, "files")
# FAISS_INDEX_PATH = "faiss_index"
# METADATA_PATH = "file_metadata.json"

# # =========================
# # CLIENTS
# # =========================
# cohere_client = cohere.ClientV2(
#     api_key=os.getenv("COHERE_API_KEY")
# )

# groq_client = Groq(
#     api_key=os.getenv("GROQ_API_KEY")
# )

# # =========================
# # FILE HASH (SMART SKIP)
# # =========================
# def get_file_hash(path):
#     hasher = hashlib.md5()

#     with open(path, "rb") as f:
#         while chunk := f.read(8192):
#             hasher.update(chunk)

#     return hasher.hexdigest()


# # =========================
# # EMBEDDINGS
# # =========================
# class CohereEmbeddings(Embeddings):

#     def __init__(self):
#         self.model = "embed-v4.0"
#         self.cache = {}

#     def _safe_embed(self, texts, input_type):

#         inputs = []

#         for text in texts:
#             inputs.append(
#                 {
#                     "content": [
#                         {
#                             "type": "text",
#                             "text": text
#                         }
#                     ]
#                 }
#             )

#         for attempt in range(10):
#             try:
#                 response = cohere_client.embed(
#                     inputs=inputs,
#                     model=self.model,
#                     input_type=input_type,
#                     embedding_types=["float"]
#                 )

#                 return response.embeddings.float

#             except Exception as e:
#                 print(f"⚠️ Rate limit hit (attempt {attempt+1})... retrying")
#                 time.sleep(15 + attempt * 10)

#         raise Exception("❌ Embedding failed after retries")

#     def embed_documents(self, texts: List[str]) -> List[List[float]]:

#         batch_size = 3
#         results = []

#         for i in range(0, len(texts), batch_size):
#             batch = texts[i:i + batch_size]

#             embeddings = self._safe_embed(
#                 batch,
#                 "search_document"
#             )

#             results.extend(embeddings)

#             time.sleep(8)

#         return results

#     def embed_query(self, text: str) -> List[float]:

#         if text in self.cache:
#             return self.cache[text]

#         embedding = self._safe_embed(
#             [text],
#             "search_query"
#         )[0]

#         self.cache[text] = embedding

#         return embedding


# # =========================
# # GROQ OCR
# # =========================
# def groq_image_ocr(image_bytes):

#     b64 = base64.b64encode(image_bytes).decode()

#     response = groq_client.chat.completions.create(
#         model="meta-llama/llama-4-scout-17b-16e-instruct",
#         messages=[
#             {
#                 "role": "user",
#                 "content": [
#                     {
#                         "type": "text",
#                         "text": "Extract ALL text from this document exactly."
#                     },
#                     {
#                         "type": "image_url",
#                         "image_url": {
#                             "url": f"data:image/png;base64,{b64}"
#                         }
#                     }
#                 ]
#             }
#         ]
#     )

#     return response.choices[0].message.content


# # =========================
# # DOCUMENT LOADER
# # =========================
# class DocumentLoader:

#     def __init__(self):

#         self.splitter = RecursiveCharacterTextSplitter(
#             chunk_size=700,
#             chunk_overlap=100
#         )

#         self.embeddings = CohereEmbeddings()

#     def load_from_folder(self, folder_path=DOCUMENTS_FOLDER):

#         os.makedirs(folder_path, exist_ok=True)

#         files = list(
#             set(
#                 glob.glob(
#                     os.path.join(folder_path, "**", "*.pdf"),
#                     recursive=True
#                 )
#                 +
#                 glob.glob(
#                     os.path.join(folder_path, "**", "*.docx"),
#                     recursive=True
#                 )
#             )
#         )

#         if not files:
#             raise FileNotFoundError("❌ No documents found")

#         # Load metadata
#         old_metadata = {}

#         if os.path.exists(METADATA_PATH):
#             with open(METADATA_PATH, "r") as f:
#                 old_metadata = json.load(f)

#         vectorstore = None

#         # Load existing FAISS
#         if os.path.exists(FAISS_INDEX_PATH):
#             try:
#                 vectorstore = FAISS.load_local(
#                     FAISS_INDEX_PATH,
#                     self.embeddings,
#                     allow_dangerous_deserialization=True
#                 )
#                 print("✅ Loaded FAISS index")

#             except:
#                 vectorstore = None
#                 print("⚠️ Rebuilding index...")

#         new_docs = []
#         updated_metadata = old_metadata.copy()

#         print(f"\n📂 Total Files: {len(files)}")

#         for fp in files:

#             fname = os.path.basename(fp)
#             file_hash = get_file_hash(fp)

#             # Skip unchanged files
#             if fname in old_metadata and old_metadata[fname] == file_hash:
#                 print(f"⏩ Skipped: {fname}")
#                 continue

#             print(f"\n🔍 Processing: {fname}")

#             text = self._extract_text(fp)

#             if not text.strip():
#                 print("❌ Empty file")
#                 continue

#             chunks = self.splitter.split_text(text)

#             docs = [
#                 Document(
#                     page_content=c,
#                     metadata={"source": fname}
#                 )
#                 for c in chunks
#             ]

#             new_docs.extend(docs)
#             updated_metadata[fname] = file_hash

#             print(f"✅ Chunks: {len(chunks)}")

#         # Build / Update FAISS
#         if vectorstore is None:

#             if not new_docs:
#                 raise ValueError("❌ No documents to index")

#             vectorstore = FAISS.from_documents(
#                 new_docs,
#                 self.embeddings
#             )

#             print("\n🆕 Created new FAISS index")

#         elif new_docs:

#             vectorstore.add_documents(new_docs)

#             print("\n➕ Added new documents")

#         else:
#             print("\n⚡ No new documents — skipped embedding")

#         # Save
#         vectorstore.save_local(FAISS_INDEX_PATH)

#         with open(METADATA_PATH, "w") as f:
#             json.dump(updated_metadata, f)

#         return vectorstore

#     # =========================
#     # EXTRACTION
#     # =========================
#     def _extract_text(self, path):

#         if path.endswith(".pdf"):
#             return self._extract_pdf(path)

#         return self._extract_docx(path)

#     def _extract_pdf(self, path):

#         # FAST (PyMuPDF)
#         try:
#             import fitz

#             text = ""

#             with fitz.open(path) as doc:
#                 for p in doc:
#                     text += p.get_text()

#             if text.strip():
#                 print("📄 PyMuPDF success")
#                 return text

#         except:
#             pass

#         # GROQ OCR
#         print("🚀 Using Groq OCR")

#         try:
#             import fitz

#             doc = fitz.open(path)

#             final_text = ""

#             for page in doc:
#                 pix = page.get_pixmap()
#                 img_bytes = pix.tobytes("png")

#                 page_text = groq_image_ocr(img_bytes)

#                 final_text += page_text + "\n\n"

#             return final_text

#         except Exception as e:
#             print("❌ Groq OCR failed:", e)

#         return ""

#     def _extract_docx(self, path):

#         try:
#             from docx import Document as DocxDoc

#             doc = DocxDoc(path)

#             text_parts = []

#             for p in doc.paragraphs:
#                 if p.text.strip():
#                     text_parts.append(p.text)

#             for table in doc.tables:
#                 for row in table.rows:
#                     for cell in row.cells:
#                         if cell.text.strip():
#                             text_parts.append(cell.text)

#             return "\n\n".join(text_parts)

#         except Exception as e:
#             print("❌ DOCX failed:", e)
#             return ""





















# # ## from groq api embedding 

# # import os
# # import glob
# # import json
# # import time
# # import base64
# # import hashlib
# # from typing import List

# # # import cohere
# # from groq import Groq

# # from langchain_core.embeddings import Embeddings
# # from langchain_text_splitters import RecursiveCharacterTextSplitter
# # from langchain_core.documents import Document
# # from langchain_community.vectorstores import FAISS

# # # =========================
# # # PATHS
# # # =========================
# # BASE_DIR = os.path.dirname(os.path.dirname(__file__))
# # DOCUMENTS_FOLDER = os.path.join(BASE_DIR, "files")
# # FAISS_INDEX_PATH = "faiss_index"
# # METADATA_PATH = "file_metadata.json"

# # # =========================
# # # CLIENTS
# # # =========================
# # cohere_client = cohere.Client(os.getenv("COHERE_API_KEY"))
# # groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# # # =========================
# # # FILE HASH (SMART SKIP)
# # # =========================
# # def get_file_hash(path):
# #     hasher = hashlib.md5()
# #     with open(path, "rb") as f:
# #         while chunk := f.read(8192):
# #             hasher.update(chunk)
# #     return hasher.hexdigest()

# # # =========================
# # # EMBEDDINGS
# # # =========================
# # class CohereEmbeddings(Embeddings):

# #     def __init__(self):
# #         self.model = "embed-v4.0"                     #"embed-english-v3.0"
# #         self.cache = {}

# #     def _embed(self, texts, input_type):
# #         return cohere_client.embed(
# #             texts=texts,
# #             model=self.model,
# #             input_type=input_type
# #         )

# #     def embed_documents(self, texts: List[str]) -> List[List[float]]:
# #         batch_size = 20
# #         results = []

# #         for i in range(0, len(texts), batch_size):
# #             batch = texts[i:i+batch_size]
# #             res = self._embed(batch, "search_document")
# #             results.extend(res.embeddings)
# #             time.sleep(1)

# #         return results

# #     def embed_query(self, text: str) -> List[float]:
# #         if text in self.cache:
# #             return self.cache[text]

# #         res = self._embed([text], "search_query")
# #         emb = res.embeddings[0]
# #         self.cache[text] = emb
# #         return emb

# # # =========================
# # # GROQ OCR
# # # =========================
# # def groq_image_ocr(image_bytes):

# #     b64 = base64.b64encode(image_bytes).decode()

# #     response = groq_client.chat.completions.create(
# #         model="meta-llama/llama-4-scout-17b-16e-instruct",
# #         messages=[
# #             {
# #                 "role": "user",
# #                 "content": [
# #                     {"type": "text", "text": "Extract ALL text from this document exactly."},
# #                     {
# #                         "type": "image_url",
# #                         "image_url": {
# #                             "url": f"data:image/png;base64,{b64}"
# #                         }
# #                     }
# #                 ]
# #             }
# #         ]
# #     )

# #     return response.choices[0].message.content

# # # =========================
# # # DOCUMENT LOADER
# # # =========================
# # class DocumentLoader:

# #     def __init__(self):
# #         self.splitter = RecursiveCharacterTextSplitter(
# #             chunk_size=700,
# #             chunk_overlap=100
# #         )
# #         self.embeddings = CohereEmbeddings()

# #     def load_from_folder(self, folder_path=DOCUMENTS_FOLDER):

# #         os.makedirs(folder_path, exist_ok=True)

# #         files = list(set(
# #             glob.glob(os.path.join(folder_path, "**", "*.pdf"), recursive=True) +
# #             glob.glob(os.path.join(folder_path, "**", "*.docx"), recursive=True)
# #         ))

# #         if not files:
# #             raise FileNotFoundError("❌ No documents found")

# #         # Load metadata
# #         old_metadata = {}
# #         if os.path.exists(METADATA_PATH):
# #             with open(METADATA_PATH, "r") as f:
# #                 old_metadata = json.load(f)

# #         vectorstore = None

# #         # Load existing FAISS
# #         if os.path.exists(FAISS_INDEX_PATH):
# #             try:
# #                 vectorstore = FAISS.load_local(
# #                     FAISS_INDEX_PATH,
# #                     self.embeddings,
# #                     allow_dangerous_deserialization=True
# #                 )
# #                 print("✅ Loaded FAISS index")
# #             except:
# #                 vectorstore = None
# #                 print("⚠️ Rebuilding index...")

# #         new_docs = []
# #         updated_metadata = old_metadata.copy()

# #         print(f"\n📂 Total Files: {len(files)}")

# #         for fp in files:
# #             fname = os.path.basename(fp)
# #             file_hash = get_file_hash(fp)

# #             # Skip unchanged files
# #             if fname in old_metadata and old_metadata[fname] == file_hash:
# #                 print(f"⏩ Skipped: {fname}")
# #                 continue

# #             print(f"\n🔍 Processing: {fname}")

# #             text = self._extract_text(fp)

# #             if not text.strip():
# #                 print("❌ Empty file")
# #                 continue

# #             chunks = self.splitter.split_text(text)

# #             docs = [
# #                 Document(page_content=c, metadata={"source": fname})
# #                 for c in chunks
# #             ]

# #             new_docs.extend(docs)
# #             updated_metadata[fname] = file_hash

# #             print(f"✅ Chunks: {len(chunks)}")

# #         # Build / Update FAISS
# #         if vectorstore is None:
# #             if not new_docs:
# #                 raise ValueError("❌ No documents to index")

# #             vectorstore = FAISS.from_documents(new_docs, self.embeddings)
# #             print("\n🆕 Created new FAISS index")

# #         elif new_docs:
# #             vectorstore.add_documents(new_docs)
# #             print("\n➕ Added new documents")

# #         else:
# #             print("\n⚡ No new documents — skipped embedding")

# #         # Save
# #         vectorstore.save_local(FAISS_INDEX_PATH)

# #         with open(METADATA_PATH, "w") as f:
# #             json.dump(updated_metadata, f)

# #         return vectorstore

# #     # =========================
# #     # EXTRACTION
# #     # =========================
# #     def _extract_text(self, path):

# #         if path.endswith(".pdf"):
# #             return self._extract_pdf(path)

# #         return self._extract_docx(path)

# #     def _extract_pdf(self, path):

# #         # FAST (PyMuPDF)
# #         try:
# #             import fitz
# #             text = ""

# #             with fitz.open(path) as doc:
# #                 for p in doc:
# #                     text += p.get_text()

# #             if text.strip():
# #                 print("📄 PyMuPDF success")
# #                 return text

# #         except:
# #             pass

# #         # GROQ OCR
# #         print("🚀 Using Groq OCR")

# #         try:
# #             import fitz
# #             doc = fitz.open(path)

# #             final_text = ""

# #             for page in doc:
# #                 pix = page.get_pixmap()
# #                 img_bytes = pix.tobytes("png")

# #                 page_text = groq_image_ocr(img_bytes)
# #                 final_text += page_text + "\n\n"

# #             return final_text

# #         except Exception as e:
# #             print("❌ Groq OCR failed:", e)

# #         return ""

# #     def _extract_docx(self, path):

# #         try:
# #             from docx import Document as DocxDoc

# #             doc = DocxDoc(path)
# #             text_parts = []

# #             for p in doc.paragraphs:
# #                 if p.text.strip():
# #                     text_parts.append(p.text)

# #             for table in doc.tables:
# #                 for row in table.rows:
# #                     for cell in row.cells:
# #                         if cell.text.strip():
# #                             text_parts.append(cell.text)

# #             return "\n\n".join(text_parts)

# #         except Exception as e:
# #             print("❌ DOCX failed:", e)
# #             return ""